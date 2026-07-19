"""
PDF and DOCX report generation utilities.
"""

import logging
from pathlib import Path

from app.config import settings

logger = logging.getLogger(__name__)


def generate_pdf_report(content: dict, candidate_name: str) -> str:
    """
    Generate a PDF evaluation report.

    Returns the file path to the generated PDF.
    """
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Title
        pdf.set_font("Helvetica", "B", 20)
        pdf.cell(0, 15, "Candidate Evaluation Report", ln=True, align="C")
        pdf.ln(5)

        # Candidate Info
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, f"Candidate: {candidate_name}", ln=True)
        pdf.ln(3)

        candidate_info = content.get("candidate", {})
        pdf.set_font("Helvetica", "", 10)
        if candidate_info.get("email"):
            pdf.cell(0, 6, f"Email: {candidate_info['email']}", ln=True)
        if candidate_info.get("phone"):
            pdf.cell(0, 6, f"Phone: {candidate_info['phone']}", ln=True)
        pdf.ln(5)

        # Scores section
        scores = content.get("scores", {})
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Scores", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 6, f"Resume Score: {scores.get('resume_score', 'N/A')}/100", ln=True)
        pdf.cell(0, 6, f"Interview Score: {scores.get('interview_score', 'N/A')}/100", ln=True)
        pdf.cell(0, 6, f"Final Score: {scores.get('final_score', 'N/A')}/100", ln=True)

        rec = scores.get("recommendation", "N/A")
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, f"Recommendation: {rec.replace('_', ' ').title()}", ln=True)
        pdf.ln(5)

        # Executive Summary
        _add_section(pdf, "Executive Summary", content.get("executive_summary", ""))

        # Resume Summary
        _add_section(pdf, "Resume Summary", content.get("resume_summary", ""))

        # Interview Summary
        _add_section(pdf, "Interview Summary", content.get("interview_summary", ""))

        # Strengths
        _add_section(pdf, "Combined Strengths", content.get("combined_strengths", ""))

        # Weaknesses
        _add_section(pdf, "Areas for Improvement", content.get("combined_weaknesses", ""))

        # Missing Skills
        missing = content.get("missing_skills", [])
        if missing:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "Missing Skills", ln=True)
            pdf.set_font("Helvetica", "", 10)
            for skill in missing:
                pdf.cell(0, 6, f"  - {skill}", ln=True)
            pdf.ln(3)

        # Hiring Recommendation
        _add_section(pdf, "Hiring Recommendation", content.get("hiring_recommendation", ""))

        # Save
        output_dir = settings.upload_path / "reports"
        output_dir.mkdir(parents=True, exist_ok=True)
        safe_name = candidate_name.replace(" ", "_").lower()
        file_path = str(output_dir / f"report_{safe_name}.pdf")
        pdf.output(file_path)

        logger.info(f"PDF report generated: {file_path}")
        return file_path

    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        raise ValueError(f"Failed to generate PDF report: {e}")


def _add_section(pdf, title: str, content: str):
    """Helper to add a titled section to the PDF."""
    if not content:
        return
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, title, ln=True)
    pdf.set_font("Helvetica", "", 10)
    # Use multi_cell for word-wrapped text
    pdf.multi_cell(0, 6, content)
    pdf.ln(3)


def generate_docx_report(content: dict, candidate_name: str) -> str:
    """
    Generate a DOCX evaluation report.

    Returns the file path to the generated DOCX.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading("Candidate Evaluation Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Candidate Info
        doc.add_heading(f"Candidate: {candidate_name}", level=1)
        candidate_info = content.get("candidate", {})
        if candidate_info.get("email"):
            doc.add_paragraph(f"Email: {candidate_info['email']}")
        if candidate_info.get("phone"):
            doc.add_paragraph(f"Phone: {candidate_info['phone']}")

        # Scores
        doc.add_heading("Scores", level=2)
        scores = content.get("scores", {})
        table = doc.add_table(rows=4, cols=2)
        table.style = "Light List Accent 1"
        rows_data = [
            ("Resume Score", f"{scores.get('resume_score', 'N/A')}/100"),
            ("Interview Score", f"{scores.get('interview_score', 'N/A')}/100"),
            ("Final Score", f"{scores.get('final_score', 'N/A')}/100"),
            ("Recommendation", str(scores.get('recommendation', 'N/A')).replace('_', ' ').title()),
        ]
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value

        # Sections
        sections = [
            ("Executive Summary", content.get("executive_summary", "")),
            ("Resume Summary", content.get("resume_summary", "")),
            ("Interview Summary", content.get("interview_summary", "")),
            ("Combined Strengths", content.get("combined_strengths", "")),
            ("Areas for Improvement", content.get("combined_weaknesses", "")),
            ("Hiring Recommendation", content.get("hiring_recommendation", "")),
        ]

        for title, text in sections:
            if text:
                doc.add_heading(title, level=2)
                doc.add_paragraph(text)

        # Missing Skills
        missing = content.get("missing_skills", [])
        if missing:
            doc.add_heading("Missing Skills", level=2)
            for skill in missing:
                doc.add_paragraph(skill, style="List Bullet")

        # Save
        output_dir = settings.upload_path / "reports"
        output_dir.mkdir(parents=True, exist_ok=True)
        safe_name = candidate_name.replace(" ", "_").lower()
        file_path = str(output_dir / f"report_{safe_name}.docx")
        doc.save(file_path)

        logger.info(f"DOCX report generated: {file_path}")
        return file_path

    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        raise ValueError(f"Failed to generate DOCX report: {e}")
